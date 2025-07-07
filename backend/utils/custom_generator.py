import os
import sqlite3
import pandas as pd
import logging
from django.conf import settings
from mis.settings import MASTER_DATA, STORAGE_PATH, COUNTRY_NAME
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from api.v1.v1_profile.models import Administration

logger = logging.getLogger(__name__)


def generate_sqlite(model, test: bool = False):
    if not test:
        test = settings.TEST_ENV
    table_name = model._meta.db_table
    field_names = [f.name for f in model._meta.fields]
    objects = model.objects.all()
    file_name = "{0}/{1}{2}.sqlite".format(
        MASTER_DATA,
        "test_" if test else "",
        table_name,
    )
    if os.path.exists(file_name):
        os.remove(file_name)
    data = pd.DataFrame(list(objects.values(*field_names)))
    no_rows = data.shape[0]
    if no_rows < 1:
        return
    # Add full_path_name for Administration model
    if model.__name__ == "Administration":
        # Get all Administration objects with their full_path_name property
        full_path_names = {
            obj.id: obj.full_path_name.replace("|", " - ") for obj in objects
        }
        # Add the full_path_name column to the DataFrame
        data["full_path_name"] = data["id"].apply(
            lambda id_: full_path_names.get(id_, "")
        )

    if "parent" in field_names:
        data["parent"] = data["parent"].apply(
            lambda x: int(x) if x == x else 0
        )
    elif "administration" in field_names:
        data["parent"] = data["administration"].apply(
            lambda x: int(x) if x == x else 0
        )
    else:
        data["parent"] = 0
    conn = sqlite3.connect(file_name)
    data.to_sql("nodes", conn, if_exists="replace", index=False)
    conn.close()
    return file_name


def update_sqlite(model, data, id=None):
    test = settings.TEST_ENV
    table_name = model._meta.db_table
    fields = data.keys()
    field_names = ", ".join([f for f in fields])
    placeholders = ", ".join(["?" for _ in range(len(fields))])
    update_placeholders = ", ".join([f"{f} = ?" for f in fields])
    params = list(data.values())
    if id:
        params += [id]
    file_name = "{0}/{1}{2}.sqlite".format(
        MASTER_DATA,
        "test_" if test else "",
        table_name,
    )
    conn = sqlite3.connect(file_name)
    try:
        with conn:
            c = conn.cursor()
            if id:
                c.execute("SELECT * FROM nodes WHERE id = ?", (id,))
                if c.fetchone():
                    query = f"UPDATE nodes \
                        SET {update_placeholders} WHERE id = ?"
                    c.execute(query, params)
            if not id:
                query = f"INSERT INTO nodes({field_names}) \
                    VALUES ({placeholders})"
                c.execute(query, params)
    except sqlite3.OperationalError:
        generate_sqlite(model=model, test=test)
    finally:
        conn.close()


def administration_csv_add(data: dict):
    test = settings.TEST_ENV
    filename = "{0}-administration.csv".format(
        "test" if test else COUNTRY_NAME
    )
    filepath = f"{STORAGE_PATH}/master_data/{filename}"
    if os.path.exists(filepath):
        df = pd.read_csv(filepath)
        new_data = {}
        if data.path:
            parent_ids = list(filter(lambda path: path, data.path.split(".")))
            parents = Administration.objects.filter(
                pk__in=parent_ids, level__id__gt=1
            ).all()
            for p in parents:
                new_data[p.level.name.lower()] = p.name
                new_data[f"{p.level.name.lower()}_id"] = p.id
        new_data[data.level.name.lower()] = data.name
        new_data[f"{data.level.name.lower()}_id"] = data.id
        new_df = pd.DataFrame([new_data])
        df = pd.concat([df, new_df], ignore_index=True)
        df.to_csv(filepath, index=False)
        return filepath
    else:
        logger.error(
            {
                "context": "insert_administration_row_csv",
                "message": (
                    f"{('test' if test else COUNTRY_NAME)}-administration.csv"
                    " doesn't exist"
                ),
            }
        )  # pragma: no cover
    return None


def find_index_by_id(df, id):
    for idx, row in df.iterrows():
        last_non_null_col = row.last_valid_index()
        last_non_null_value = row[last_non_null_col]
        if last_non_null_value == id:
            return idx
    return None


def administration_csv_update(data: dict):
    test = settings.TEST_ENV
    filename = "{0}-administration.csv".format(
        "test" if test else COUNTRY_NAME
    )
    filepath = f"{STORAGE_PATH}/master_data/{filename}"
    if os.path.exists(filepath):
        df = pd.read_csv(filepath)
        index = find_index_by_id(df=df, id=data.pk)
        if index is not None:
            if data.path:
                parent_ids = list(
                    filter(lambda path: path, data.path.split("."))
                )
                parents = Administration.objects.filter(
                    pk__in=parent_ids, level__id__gt=1
                ).all()
                for p in parents:
                    df.loc[index, p.level.name.lower()] = p.name
                    df.loc[index, f"{p.level.name.lower()}_id"] = p.id
            df.loc[index, data.level.name.lower()] = data.name
            df.loc[index, f"{data.level.name.lower()}_id"] = data.id
        df.to_csv(filepath, index=False)
        return filepath
    else:
        logger.error(
            {
                "context": "update_administration_row_csv",
                "message": f"{filename} doesn't exist",
            }
        )  # pragma: no cover
    return None


def administration_csv_delete(id: int):
    test = settings.TEST_ENV
    filename = "{0}-administration.csv".format(
        "test" if test else COUNTRY_NAME
    )
    filepath = f"{STORAGE_PATH}/master_data/{filename}"
    if os.path.exists(filepath):
        df = pd.read_csv(filepath)
        ix = find_index_by_id(df=df, id=id)
        if ix is not None:
            df.drop(index=ix, inplace=True)
        df.to_csv(filepath, index=False)
        return filepath
    else:
        logger.error(
            {
                "context": "delete_administration_row_csv",
                "message": f"{filename} doesn't exist",
            }
        )  # pragma: no cover
    return None


def generate_datapoint_report(
    report_data, file_path="./tmp/inspection_report.docx"
):
    """
    Generates a .docx report for a single inspection record.

    Args:
        report_data (dict):
        A dictionary containing all the data for one report.
        file_path (str): The full path where the document should be saved.
    """

    # --- Document Initialization ---
    document = Document()

    # Set default font for the document (optional)
    style = document.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)

    # --- Helper Function to add key-value pairs ---
    def add_info(key, value):
        """Adds a formatted key-value pair to the document."""
        p = document.add_paragraph()
        p.add_run(f"{key}: ").bold = True
        p.add_run(str(value))

    # --- Header ---
    title = document.add_heading(
        "EPS Inspection and Water Quality Monitoring", level=1
    )
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Add a subtitle for the specific village
    subtitle = document.add_heading(
        f"Report for: {report_data.get('Village Name', 'N/A')}",
        level=2,
    )
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.add_paragraph()  # Add some space

    # --- General Information ---
    document.add_heading("General Information", level=3)
    add_info("Display Name", report_data.get("Display Name"))
    add_info("Identifier", report_data.get("Identifier"))
    add_info("Device Identifier", report_data.get("Device identifier"))
    add_info("Submission Date", report_data.get("Submission Date"))
    add_info("Submitter", report_data.get("Submitter"))
    add_info("Form Version", report_data.get("Form version"))
    document.add_paragraph()

    # --- Group 1: Location ---
    document.add_heading("Group 1 - Location of the Village", level=3)
    location_data = report_data.get("Group 1", {})
    add_info(
        "Division-Province-Tikina",
        location_data.get("Which Division-Province-Tikina are you in?"),
    )
    add_info("Village Name", location_data.get("Village Name"))
    add_info("Date of Inspection", location_data.get("Date of Inspection"))
    add_info("Weather Condition", location_data.get("Weather Condition"))
    add_info("Water Source Type", location_data.get("Type of Water Source"))
    add_info(
        "Project Implementation Date",
        location_data.get("Project Implementation date"),
    )
    add_info(
        "Coordinates (Lat, Lon)",
        f"{location_data.get('Latitude', 'N/A')}, "
        f"{location_data.get('Longitude', 'N/A')}",
    )
    add_info("Elevation", location_data.get("Elevation"))
    document.add_paragraph()

    # --- Group 2: Contacts ---
    document.add_heading("Group 2 - Contact Details", level=3)
    contact_data = report_data.get("Group 2", {})
    add_info(
        "Village Headman/TNK/Nurse Name",
        contact_data.get(
            "Name of the village headman/TNK/village nurse?"
        ),
    )
    add_info(
        "Phone Contact",
        contact_data.get(
            "Phone contact of the Village Headman/TNK or village nurse?"
        ),
    )
    document.add_paragraph()

    # --- Group 3: Photos ---
    document.add_heading("Group 3 - EPS Photos", level=3)
    photo_data = report_data.get("Group 3", {})
    add_info("Comments on Photo", photo_data.get("Comment on Photo taken"))

    # Placeholder for adding an image.
    # To add a real image,
    # uncomment the following lines and provide a valid path.
    # try:
    #     document.add_picture('path/to/your/image.jpg', width=Inches(4.0))
    # except FileNotFoundError:
    #     document.add_paragraph(
    #   "Image file not found. Please provide a valid path."
    # )
    document.add_paragraph()

    # --- Group 4: Water Quality ---
    document.add_heading("Group 4 - Water Quality Testing", level=3)
    quality_data = report_data.get("Group 4", {})
    add_info(
        "Method of Water Testing",
        quality_data.get("Method of Water Testing Used?"),
    )
    add_info(
        "Description of Sampling Point",
        quality_data.get("Description of Sampling Point"),
    )
    add_info(
        "Health Risk Category",
        quality_data.get(
            "Health Risk Category (Based on MPN and Confidence Interval)"
        ),
    )
    add_info("MPN (MPN/100ml)", quality_data.get("MPN(MPN/100ml)"))
    add_info(
        "Upper 95% Confidence Interval",
        quality_data.get("Upper 95% Confidence Interval"),
    )
    document.add_paragraph()

    # --- Final Remarks ---
    document.add_heading("Concluding Remarks", level=3)
    add_info("General Remarks", report_data.get("General Remarks"))
    add_info(
        "Current System Status",
        report_data.get("The current status of this system?"),
    )
    add_info("Signature of Officer", report_data.get("Signature of Officer"))

    # --- Save the document ---
    try:
        # Ensure the directory exists
        os.makedirs(os.path.dirname(file_path), exist_ok=True)
        document.save(file_path)
        return file_path
    except Exception as e:
        logger.error(
            {
                "context": "generate_datapoint_report",
                "message": e,
            }
        )  # pragma: no cover
        raise
