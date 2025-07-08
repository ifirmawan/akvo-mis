import os
import base64
import tempfile
import logging
from mis.settings import STORAGE_PATH
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

logger = logging.getLogger(__name__)


def generate_datapoint_report(
    report_data: list,
    file_path: str = "./tmp/inspection_report.docx",
    form_name: str = "EPS Inspection and Water Quality Monitoring",
    display_names: list = []
):
    """
    Generates a .docx report with multiple answer columns across multiple
    tables.

    Args:
        report_data (list):
        A list of groups, each containing a name and questions array.
        Each group has the structure:
        {
            "name": "Group Name",
            "questions": [
                {
                    "question": "Question text",
                    "answers": ["Answer value 1", "Answer value 2", ...]
                }
            ]
        }

        All groups will be distributed across multiple tables with group
        headers. Multiple answers will be displayed in separate columns
        rather than concatenated with "|" separators. When more than 5
        answers exist, the data will be split into multiple tables (batches)
        with page breaks.

        For photo questions, answers can contain image file paths:
        {
            "question": "Photo of Water Quality Test",
            "answers": ["/images/photo1.jpg", "/images/photo2.png"]
        }
        Each image will be rendered in its own column across the appropriate
        tables.

        file_path (str): The full path where the document should be saved.
    """

    # --- Document Initialization ---
    document = Document()

    # Set page orientation to landscape and margins
    section = document.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Inches(11)
    section.page_height = Inches(8.5)
    # Set 1-inch margins on all sides
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)

    # Set default font for the document (optional)
    style = document.styles["Normal"]
    font = style.font
    font.name = "Times New Roman"
    font.size = Pt(12)

    # --- Header ---
    title = document.add_heading(form_name, level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Extract village name from the first group's questions for subtitle
    datapoint_name = "N/A"
    for group in report_data:
        for question_data in group.get("questions", []):
            if question_data.get("question", "") == "Village Name":
                answers = question_data.get("answers", [])
                if answers:
                    # Use the first village name if multiple, or join them
                    if len(answers) == 1:
                        datapoint_name = str(answers[0])
                    else:
                        datapoint_name = " & ".join(
                            str(ans) for ans in answers
                        )
                break
        if datapoint_name != "N/A":
            break

    # --- Calculate overall max answers across all groups ---
    overall_max_answers = 1  # At least 1 column for answers
    for group in report_data:
        questions = group.get("questions", [])
        for question_data in questions:
            answers = question_data.get("answers", [])
            # Skip coordinate questions as they're handled specially
            question_text = question_data.get("question", "")
            if question_text not in ["Latitude", "Longitude"]:
                overall_max_answers = max(overall_max_answers, len(answers))

    # Determine how many tables we need based on max answers
    max_answers_per_table = 5
    tables_needed = 1
    if overall_max_answers > max_answers_per_table:
        tables_needed = (
            overall_max_answers + max_answers_per_table - 1
        ) // max_answers_per_table

    # Create tables as needed
    tables = []
    for table_idx in range(tables_needed):
        start_answer_idx = table_idx * max_answers_per_table
        end_answer_idx = min(
            (table_idx + 1) * max_answers_per_table, overall_max_answers
        )
        answers_in_table = max(1, end_answer_idx - start_answer_idx)

        # Skip tables that would have no valid range
        if start_answer_idx >= overall_max_answers:
            continue

        # 1 column for question + answers_in_table columns for answers
        tables.append(
            {
                "start_idx": start_answer_idx,
                "end_idx": end_answer_idx,
                "total_cols": 1 + answers_in_table,
            }
        )

    # --- Process Each Table Separately (Batch Processing) ---
    for table_info in tables:
        start_idx = table_info["start_idx"]
        end_idx = table_info["end_idx"]
        total_cols = table_info["total_cols"]

        table = document.add_table(rows=0, cols=total_cols)
        table.style = "Table Grid"

        header_row = table.add_row()
        header_row.cells[0].text = "Identifier"
        for i in range(1, total_cols):
            # Use display names if provided, otherwise default to "Data #i"
            header_row.cells[i].text = f"Data #{i}"
            if display_names[i - 1]:
                header_row.cells[i].text = display_names[i - 1]

        for cell in header_row.cells:
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.bold = True
                    run.font.size = Pt(12)
            # Set background color for header row
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            shd = OxmlElement('w:shd')
            shd.set(qn('w:fill'), 'a8d08d')
            tcPr.append(shd)
        # Set header row as repeating header
        tbl_header = OxmlElement('w:tblHeader')
        header_row._tr.get_or_add_trPr().append(tbl_header)

        # Process each group for this specific table
        for group in report_data:
            group_name = group.get("name", "Unknown Group")
            questions = group.get("questions", [])

            # Add group header row for this table
            group_header_row = table.add_row()
            if len(group_header_row.cells) > 0:
                merged_cell = group_header_row.cells[0]
                for i in range(
                    1, min(total_cols, len(group_header_row.cells))
                ):
                    merged_cell = merged_cell.merge(group_header_row.cells[i])
                merged_cell.text = group_name
                # Make header text bold and larger
                for paragraph in merged_cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in paragraph.runs:
                        run.bold = True
                        run.font.size = Pt(12)

            # Add all questions for this group in this table's batch
            for question_data in questions:
                question = question_data.get("question", "")
                answers = question_data.get("answers", [])

                # Handle special formatting for coordinates
                if question == "Latitude" and "Longitude" in [
                    q.get("question", "") for q in questions
                ]:
                    # Find longitude answers
                    longitude_data = next(
                        (
                            q
                            for q in questions
                            if q.get("question", "") == "Longitude"
                        ),
                        {},
                    )
                    longitude_answers = longitude_data.get("answers", [])

                    # Combine latitude and longitude for each pair
                    coord_pairs = []
                    max_len = max(len(answers), len(longitude_answers))
                    for i in range(max_len):
                        lat = answers[i] if i < len(answers) else ""
                        lon = (
                            longitude_answers[i]
                            if i < len(longitude_answers)
                            else ""
                        )
                        coord_pairs.append(f"{lat}, {lon}")

                    # Add coordinate row for this table's batch
                    coord_row = table.add_row()
                    coord_row.cells[0].text = "Coordinates (Lat, Lon)"
                    # Make question cell bold
                    for paragraph in coord_row.cells[0].paragraphs:
                        for run in paragraph.runs:
                            run.bold = True

                    # Fill coordinate pairs for this table's range (batch)
                    for i in range(start_idx, end_idx):
                        col_idx = i - start_idx + 1
                        if i < len(coord_pairs) and col_idx < total_cols:
                            safe_set_cell_text(
                                coord_row, col_idx, coord_pairs[i]
                            )

                    # Fill remaining cells with empty content
                    remaining_start = max(
                        1,
                        min(len(coord_pairs) - start_idx, end_idx - start_idx)
                        + 1,
                    )
                    for col_idx in range(remaining_start, total_cols):
                        safe_set_cell_text(coord_row, col_idx, "")

                    continue
                elif question == "Longitude":
                    # Skip this as it's handled with Latitude
                    continue
                else:
                    # Also check if answers contain image paths
                    has_images = any(
                        is_image_path(str(ans)) for ans in answers
                    )

                    if has_images:
                        # Always preserve the index of answers for images
                        row = table.add_row()
                        row.cells[0].text = question
                        # Make question cell bold
                        for paragraph in row.cells[0].paragraphs:
                            for run in paragraph.runs:
                                run.bold = True

                        # Add images or empty for this table's batch
                        for i in range(start_idx, end_idx):
                            col_idx = i - start_idx + 1
                            if (
                                i < len(answers) and
                                col_idx < total_cols and
                                col_idx < len(row.cells)
                            ):
                                ans = answers[i]
                                if ans and is_image_path(str(ans)):
                                    add_images_to_cell(
                                        row.cells[col_idx],
                                        [str(ans)],
                                        table=table,
                                        cell_index=col_idx,
                                    )
                                else:
                                    # Not an image or empty, leave cell empty
                                    safe_set_cell_text(row, col_idx, "")

                        # Fill remaining answer columns with empty content
                        remaining_start = max(
                            1,
                            min(len(answers) - start_idx, end_idx - start_idx)
                            + 1,
                        )
                        for col_idx in range(remaining_start, total_cols):
                            safe_set_cell_text(row, col_idx, "")
                    else:
                        # Regular questions - add answers for this batch
                        row = table.add_row()
                        if len(row.cells) > 0:
                            row.cells[0].text = question
                            # Make question cell bold
                            for paragraph in row.cells[0].paragraphs:
                                for run in paragraph.runs:
                                    run.bold = True
                        else:
                            print(
                                f"ERROR: Row has no cells! "
                                f"Table cols: {total_cols}"
                            )
                            continue

                        # Add answers for this table's batch
                        # (start_idx to end_idx)
                        for i in range(start_idx, end_idx):
                            col_idx = i - start_idx + 1
                            if i < len(answers) and col_idx < total_cols:
                                # Always write the value, even if empty string
                                safe_set_cell_text(
                                    row,
                                    col_idx,
                                    "" if answers[i] is None
                                    else str(answers[i])
                                )

                        # Fill remaining cells with empty content
                        remaining_start = max(
                            1,
                            min(len(answers) - start_idx, end_idx - start_idx)
                            + 1,
                        )
                        for col_idx in range(remaining_start, total_cols):
                            safe_set_cell_text(row, col_idx, "")

        document.add_page_break()
    # --- Save the document ---
    try:
        # Ensure the directory exists
        os.makedirs(os.path.dirname(file_path), exist_ok=True)
        document.save(file_path)
        return file_path
    except Exception as e:
        logger.error(
            {
                "context": "generate_datapoint_report",
                "message": e,
            }
        )  # pragma: no cover
        raise


def safe_set_cell_text(row, col_idx, text):
    """Safely set cell text with bounds checking."""
    try:
        if col_idx < len(row.cells):
            row.cells[col_idx].text = str(text) if text is not None else ""
            return True
        return False
    except Exception as e:
        # Debug output for troubleshooting
        cell_count = (
            len(row.cells) if hasattr(row, "cells") else "no cells attr"
        )
        print(
            f"Error in safe_set_cell_text: col_idx={col_idx}, "
            f"row.cells length={cell_count}, error={e}"
        )
        return False


def is_image_path(path_str):
    """Check if a string represents an image file path."""
    if not isinstance(path_str, str):
        return False
    image_extensions = [".jpg", ".jpeg", ".png", ".gif", ".bmp", ".tiff"]
    # or start with "data:image/"
    if path_str.startswith("data:image/"):
        return True
    # or start with http(s)://
    if path_str.startswith("http://") or path_str.startswith("https://"):
        return True
    return any(path_str.lower().endswith(ext) for ext in image_extensions)


def get_full_image_path(relative_path):
    """Convert relative image path to full file system path."""
    # Remove leading slash if present
    if relative_path.startswith("/"):
        relative_path = relative_path[1:]

    # Construct full path (assuming images are stored in storage folder)
    full_path = os.path.join(STORAGE_PATH, relative_path)
    return full_path


def create_temp_image_from_base64(base64_data, filename_prefix="temp_image"):
    """
    Create a temporary image file from base64 data.

    Args:
        base64_data (str): Base64 encoded image data
        filename_prefix (str): Prefix for the temporary filename

    Returns:
        str: Path to the temporary image file, or None if failed
    """
    try:
        # Handle data URL format
        if base64_data.startswith("data:image/"):
            # Extract the image format and base64 data
            header, base64_string = base64_data.split(",", 1)
            # Extract image format from header (e.g., "data:image/png;base64")
            image_format = header.split("/")[1].split(";")[0]
        else:
            # Assume it's a plain base64 string (default to PNG)
            base64_string = base64_data
            image_format = "png"

        # Decode base64 data
        image_bytes = base64.b64decode(base64_string)

        # Create temporary file
        with tempfile.NamedTemporaryFile(
            delete=False,
            suffix=f".{image_format}",
            prefix=f"{filename_prefix}_",
        ) as temp_file:
            temp_file.write(image_bytes)
            return temp_file.name

    except Exception as e:
        logger.warning(f"Failed to create temporary image from base64: {e}")
        return None


def get_image_path_or_create_temp(image_data):
    """
    Get the image path, creating a temporary file if it's base64 data.

    Args:
        image_data (str): Either a file path or base64 encoded image

    Returns:
        tuple: (image_path, is_temp_file) where is_temp_file indicates
               if the path points to a temporary file that should be cleaned up
    """
    if image_data.startswith("data:image/"):
        temp_path = create_temp_image_from_base64(image_data)
        return temp_path, True
    else:
        # It's a file path, get the full path
        full_path = get_full_image_path(image_data)
        return full_path, False


def add_image_to_table(table, key, image_paths, max_image_width=Inches(2.5)):
    """Add images to a table row with proper formatting."""
    row = table.add_row()
    key_cell = row.cells[0]
    value_cell = row.cells[1]
    key_cell.text = str(key)
    # Make key cell bold
    for paragraph in key_cell.paragraphs:
        for run in paragraph.runs:
            run.bold = True

    # Use the enhanced add_images_to_cell function
    add_images_to_cell(value_cell, image_paths, max_image_width, table, 1)


def calculate_optimal_image_width(table, cell_index):
    """
    Calculate the optimal image width based on available cell space.

    Args:
        table: The docx table object
        cell_index: The index of the cell in the row

    Returns:
        Optimal width in Inches for images in this cell
    """
    try:
        # Page width (landscape): 11 inches
        # Margins: 1 inch left + 1 inch right = 2 inches
        # Available content width: 11 - 2 = 9 inches
        available_width = Inches(9)

        # Get number of columns in the table
        if table.rows:
            num_columns = len(table.rows[0].cells)
        else:
            num_columns = 1

        # Calculate approximate cell width
        # Account for table borders and cell padding
        # (estimate ~0.2 inches total)
        border_padding = Inches(0.2)
        cell_width_emu = (available_width - border_padding) / num_columns

        # Use 85% of cell width to ensure images fit comfortably
        # with some padding around them
        # Convert back to inches for the calculation
        cell_width_inches = cell_width_emu / 914400  # EMU to inches conversion
        optimal_width = Inches(cell_width_inches * 0.85)

        # Set reasonable bounds: minimum 1 inch, maximum 3 inches
        min_width = Inches(1)
        max_width = Inches(3)

        return max(min_width, min(max_width, optimal_width))
    except Exception as e:
        logger.warning(f"Failed to calculate optimal image width: {e}")
        # Fallback to a conservative default
        return Inches(1.5)


def add_images_to_cell(
    cell, image_paths, max_image_width=None, table=None, cell_index=0
):
    """
    Add images to a single table cell with proper formatting and auto-sizing.
    Args:
        cell: The table cell to add images to
        image_paths: List of image paths (can be file paths or base64 strings)
        max_image_width: Optional fixed width for images
                        (if None, will calculate optimal width)
        table: The table object (used for width calculation
                        if max_image_width is None)
        cell_index: The cell index in the row (used for width calculation)
    """
    # Clear the cell and add images
    cell.text = ""
    temp_files_to_cleanup = []

    # Calculate optimal image width if not provided
    if max_image_width is None and table is not None:
        max_image_width = calculate_optimal_image_width(table, cell_index)
    elif max_image_width is None:
        # Fallback if no table provided
        max_image_width = Inches(2)

    try:
        for i, image_path in enumerate(image_paths):
            # Get the actual image path (could be file path or base64)
            actual_path, is_temp = get_image_path_or_create_temp(image_path)

            if is_temp and actual_path:
                temp_files_to_cleanup.append(actual_path)

            try:
                if actual_path and os.path.exists(actual_path):
                    # Add image to the cell
                    if i == 0:
                        paragraph = cell.paragraphs[0]
                    else:
                        paragraph = cell.add_paragraph()

                    if paragraph.runs:
                        run = paragraph.runs[0]
                    else:
                        run = paragraph.add_run()

                    run.add_picture(actual_path, width=max_image_width)

                    # Add spacing between images
                    if i < len(image_paths) - 1:
                        cell.add_paragraph()
                else:
                    # Image file not found or couldn't create from base64
                    if i == 0:
                        paragraph = cell.paragraphs[0]
                    else:
                        paragraph = cell.add_paragraph()

                    if paragraph.runs:
                        run = paragraph.runs[0]
                    else:
                        run = paragraph.add_run()

                    if image_path.startswith("data:image/"):
                        run.text = f"Error processing base64 image {i + 1}"
                    else:
                        filename = os.path.basename(image_path)
                        run.text = f"Image not found: {filename}"

            except Exception as e:
                # Error loading image
                if i == 0:
                    paragraph = cell.paragraphs[0]
                else:
                    paragraph = cell.add_paragraph()
                run = paragraph.add_run()

                if image_path.startswith("data:image/"):
                    run.text = f"Error loading base64 image {i + 1}"
                    logger.warning(f"Failed to add base64 image: {e}")
                else:
                    filename = os.path.basename(image_path)
                    run.text = f"Error loading image: {filename}"
                    logger.warning(f"Failed to add image {image_path}: {e}")

    finally:
        # Clean up temporary files
        for temp_file in temp_files_to_cleanup:
            try:
                os.unlink(temp_file)
            except Exception as e:
                logger.warning(
                    f"Failed to cleanup temporary file {temp_file}: {e}"
                )
