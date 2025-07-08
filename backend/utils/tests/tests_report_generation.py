import os
import tempfile
import unittest
from unittest.mock import patch, MagicMock
from docx import Document
from docx.shared import Inches
import base64

from utils.report_generator import (
    generate_datapoint_report,
    safe_set_cell_text,
    is_image_path,
    is_photo_question,
    get_full_image_path,
    add_images_to_cell,
    is_base64_image,
    create_temp_image_from_base64,
    get_image_path_or_create_temp
)


class TestReportGeneration(unittest.TestCase):
    """Test suite for report generation functions using static JSON data"""

    def setUp(self):
        """Set up test fixtures with static data"""
        self.temp_dir = tempfile.mkdtemp()
        self.test_file_path = os.path.join(self.temp_dir, "test_report.docx")

    def tearDown(self):
        """Clean up temporary files"""
        if os.path.exists(self.test_file_path):
            os.remove(self.test_file_path)
        os.rmdir(self.temp_dir)

    def test_basic_table_generation(self):
        """Test basic table generation with simple data structure"""
        report_data = [
            {
                "name": "Basic Information",
                "questions": [
                    {
                        "question": "Village Name",
                        "answers": ["Test Village"]
                    },
                    {
                        "question": "Population",
                        "answers": ["1500"]
                    },
                    {
                        "question": "District",
                        "answers": ["Test District"]
                    }
                ]
            }
        ]

        result_path = generate_datapoint_report(
            report_data,
            file_path=self.test_file_path,
            form_name="Basic Test Report"
        )

        # Verify file was created
        self.assertTrue(os.path.exists(result_path))
        self.assertEqual(result_path, self.test_file_path)

        # Load and verify document structure
        doc = Document(result_path)

        # Should have at least one table
        self.assertGreater(len(doc.tables), 0)

        # First table should have the basic questions
        table = doc.tables[0]
        self.assertGreater(len(table.rows), 3)  # Header + 3 questions

    def test_multi_column_answers(self):
        """Test handling of multiple answers across columns"""
        report_data = [
            {
                "name": "Water Quality Tests",
                "questions": [
                    {
                        "question": "pH Level",
                        "answers": ["7.2", "7.5", "7.1"]
                    },
                    {
                        "question": "Chlorine Level",
                        "answers": ["0.5", "0.6", "0.4"]
                    },
                    {
                        "question": "Temperature",
                        "answers": ["22°C", "23°C", "21°C"]
                    }
                ]
            }
        ]

        result_path = generate_datapoint_report(
            report_data,
            file_path=self.test_file_path,
            form_name="Multi-Column Test"
        )

        self.assertTrue(os.path.exists(result_path))

        # Load document and verify structure
        doc = Document(result_path)
        table = doc.tables[0]

        # Should have 4 columns (1 question + 3 answers)
        self.assertEqual(len(table.rows[1].cells), 4)  # Skip header row

    def test_multi_table_splitting(self):
        """Test data splitting across multiple tables when >5 answers"""
        report_data = [
            {
                "name": "Large Dataset",
                "questions": [
                    {
                        "question": "Sample ID",
                        "answers": [
                            "S001", "S002", "S003", "S004",
                            "S005", "S006", "S007"
                        ]
                    },
                    {
                        "question": "Result",
                        "answers": [
                            "Pass", "Pass", "Fail", "Pass",
                            "Pass", "Pass", "Fail"
                        ]
                    }
                ]
            }
        ]

        result_path = generate_datapoint_report(
            report_data,
            file_path=self.test_file_path,
            form_name="Multi-Table Test"
        )

        self.assertTrue(os.path.exists(result_path))

        # Load document and verify multiple tables were created
        doc = Document(result_path)
        self.assertGreater(len(doc.tables), 1)  # Should have at least 2 tables

        # First table should have 6 columns (1 question + 5 answers)
        first_table = doc.tables[0]
        self.assertEqual(len(first_table.rows[1].cells), 6)

    def test_coordinate_formatting(self):
        """Test special handling of latitude/longitude coordinates"""
        report_data = [
            {
                "name": "Location Data",
                "questions": [
                    {
                        "question": "Site Name",
                        "answers": ["Site A", "Site B"]
                    },
                    {
                        "question": "Latitude",
                        "answers": ["12.3456", "12.7890"]
                    },
                    {
                        "question": "Longitude",
                        "answers": ["78.9012", "78.5432"]
                    },
                    {
                        "question": "Elevation",
                        "answers": ["150m", "200m"]
                    }
                ]
            }
        ]

        result_path = generate_datapoint_report(
            report_data,
            file_path=self.test_file_path,
            form_name="Coordinate Test"
        )

        self.assertTrue(os.path.exists(result_path))

        # Load document and verify coordinate formatting
        doc = Document(result_path)
        table = doc.tables[0]

        # Find coordinates row
        coord_row_found = False
        for row in table.rows:
            if len(row.cells) > 0 and "Coordinates" in row.cells[0].text:
                coord_row_found = True
                # Should contain formatted lat,lon pairs
                self.assertIn("12.3456, 78.9012", row.cells[1].text)
                self.assertIn("12.7890, 78.5432", row.cells[2].text)
                break

        self.assertTrue(coord_row_found, "Coordinates row not found in table")

    @patch('utils.report_generator.os.path.exists')
    @patch('utils.report_generator.get_full_image_path')
    def test_image_handling_with_valid_paths(self, mock_get_path, mock_exists):
        """Test handling of image/photo questions with valid image paths"""
        # Mock image file existence
        mock_exists.return_value = True
        mock_get_path.side_effect = lambda x: f"/storage/{x}"

        report_data = [
            {
                "name": "Visual Documentation",
                "questions": [
                    {
                        "question": "Photo of Water Source",
                        "answers": ["/images/water1.jpg", "/images/water2.png"]
                    },
                    {
                        "question": "Site Image",
                        "answers": ["/images/site1.jpg"]
                    },
                    {
                        "question": "Description",
                        "answers": ["Clear water", "Good condition"]
                    }
                ]
            }
        ]

        with patch('utils.report_generator.add_images_to_cell') as mock_images:
            result_path = generate_datapoint_report(
                report_data,
                file_path=self.test_file_path,
                form_name="Image Test"
            )

            self.assertTrue(os.path.exists(result_path))

            # Verify images were processed
            self.assertGreater(mock_images.call_count, 0)

    def test_image_handling_with_missing_files(self):
        """Test handling of image questions when image files don't exist"""
        report_data = [
            {
                "name": "Missing Images",
                "questions": [
                    {
                        "question": "Photo of Equipment",
                        "answers": [
                            "/images/missing1.jpg", "/images/missing2.png"
                        ]
                    }
                ]
            }
        ]

        # Don't mock file existence - files won't exist
        result_path = generate_datapoint_report(
            report_data,
            file_path=self.test_file_path,
            form_name="Missing Image Test"
        )

        self.assertTrue(os.path.exists(result_path))

        # Document should still be created even with missing images
        doc = Document(result_path)
        self.assertGreater(len(doc.tables), 0)

    def test_empty_data_handling(self):
        """Test handling of empty or null data"""
        report_data = [
            {
                "name": "Empty Data Group",
                "questions": [
                    {
                        "question": "Empty Question",
                        "answers": []
                    },
                    {
                        "question": "Null Values",
                        "answers": [None, "", None]
                    },
                    {
                        "question": "Mixed Values",
                        "answers": ["Valid", None, "", "Another Valid"]
                    }
                ]
            }
        ]

        result_path = generate_datapoint_report(
            report_data,
            file_path=self.test_file_path,
            form_name="Empty Data Test"
        )

        self.assertTrue(os.path.exists(result_path))

        # Load document and verify it handles empty data gracefully
        doc = Document(result_path)
        table = doc.tables[0]

        # Should still create the table structure
        self.assertGreater(len(table.rows), 2)  # Header + questions

    def test_error_handling_invalid_path(self):
        """Test error handling when file path is invalid"""
        report_data = [
            {
                "name": "Test Group",
                "questions": [
                    {
                        "question": "Test Question",
                        "answers": ["Test Answer"]
                    }
                ]
            }
        ]

        # Try to save to a directory that doesn't exist and can't be created
        # Use a path that includes a file as a directory component
        invalid_path = "/etc/passwd/test.docx"  # passwd is a file, not dir

        # Should raise an exception
        with self.assertRaises((OSError, FileNotFoundError, PermissionError)):
            generate_datapoint_report(
                report_data,
                file_path=invalid_path,
                form_name="Error Test"
            )

    def test_village_name_extraction(self):
        """Test extraction of village name for subtitle"""
        report_data = [
            {
                "name": "Location Info",
                "questions": [
                    {
                        "question": "Village Name",
                        "answers": ["Test Village"]
                    },
                    {
                        "question": "Other Info",
                        "answers": ["Some data"]
                    }
                ]
            }
        ]

        result_path = generate_datapoint_report(
            report_data,
            file_path=self.test_file_path,
            form_name="Village Name Test"
        )

        self.assertTrue(os.path.exists(result_path))

        # Village name should be extracted (tested through document creation)
        doc = Document(result_path)
        self.assertGreater(len(doc.tables), 0)

    def test_multiple_village_names(self):
        """Test handling of multiple village names"""
        report_data = [
            {
                "name": "Multi-Village Data",
                "questions": [
                    {
                        "question": "Village Name",
                        "answers": ["Village A", "Village B", "Village C"]
                    },
                    {
                        "question": "Population",
                        "answers": ["1000", "1500", "800"]
                    }
                ]
            }
        ]

        result_path = generate_datapoint_report(
            report_data,
            file_path=self.test_file_path,
            form_name="Multi-Village Test"
        )

        self.assertTrue(os.path.exists(result_path))
        doc = Document(result_path)
        self.assertGreater(len(doc.tables), 0)

    def test_complex_multi_group_scenario(self):
        """Test complex scenario with multiple groups and mixed data types"""
        report_data = [
            {
                "name": "Basic Information",
                "questions": [
                    {
                        "question": "Village Name",
                        "answers": ["Complex Village"]
                    },
                    {
                        "question": "District",
                        "answers": ["Test District"]
                    }
                ]
            },
            {
                "name": "Location Data",
                "questions": [
                    {
                        "question": "Latitude",
                        "answers": ["12.3456", "12.7890"]
                    },
                    {
                        "question": "Longitude",
                        "answers": ["78.9012", "78.5432"]
                    }
                ]
            },
            {
                "name": "Quality Measurements",
                "questions": [
                    {
                        "question": "pH Level",
                        "answers": ["7.2", "7.5", "7.1", "7.3", "7.4", "7.0"]
                    },
                    {
                        "question": "Temperature",
                        "answers": [
                            "22°C", "23°C", "21°C", "24°C", "22°C", "23°C"
                        ]
                    }
                ]
            }
        ]

        result_path = generate_datapoint_report(
            report_data,
            file_path=self.test_file_path,
            form_name="Complex Scenario Test"
        )

        self.assertTrue(os.path.exists(result_path))

        # Load document and verify structure
        doc = Document(result_path)

        # Should have multiple tables due to >5 answers in last group
        self.assertGreater(len(doc.tables), 1)

    def test_professional_formatting(self):
        """Test professional formatting features"""
        report_data = [
            {
                "name": "Formatted Report",
                "questions": [
                    {
                        "question": "Test Question",
                        "answers": ["Test Answer"]
                    }
                ]
            }
        ]

        result_path = generate_datapoint_report(
            report_data,
            file_path=self.test_file_path,
            form_name="Formatting Test Report"
        )

        self.assertTrue(os.path.exists(result_path))

        # Load document and verify basic formatting applied
        doc = Document(result_path)

        # Should have landscape orientation
        section = doc.sections[0]
        self.assertEqual(section.page_width, Inches(11))
        self.assertEqual(section.page_height, Inches(8.5))

        # Should have proper margins
        self.assertEqual(section.top_margin, Inches(1))
        self.assertEqual(section.bottom_margin, Inches(1))
        self.assertEqual(section.left_margin, Inches(1))
        self.assertEqual(section.right_margin, Inches(1))


class TestUtilityFunctions(unittest.TestCase):
    """Test utility functions used in report generation"""

    def test_safe_set_cell_text(self):
        """Test safe_set_cell_text function"""
        # Create a mock row with cells
        mock_row = MagicMock()
        mock_cells = [MagicMock(), MagicMock(), MagicMock()]
        mock_row.cells = mock_cells

        # Test valid column index
        result = safe_set_cell_text(mock_row, 1, "test text")
        self.assertTrue(result)
        self.assertEqual(mock_cells[1].text, "test text")

        # Test invalid column index
        result = safe_set_cell_text(mock_row, 5, "test text")
        self.assertFalse(result)

        # Test None text
        result = safe_set_cell_text(mock_row, 0, None)
        self.assertTrue(result)
        self.assertEqual(mock_cells[0].text, "")

    def test_is_image_path(self):
        """Test is_image_path function"""
        # Valid image paths
        self.assertTrue(is_image_path("test.jpg"))
        self.assertTrue(is_image_path("image.PNG"))
        self.assertTrue(is_image_path("/path/to/image.jpeg"))
        self.assertTrue(is_image_path("photo.gif"))
        self.assertTrue(is_image_path("picture.bmp"))
        self.assertTrue(is_image_path("scan.tiff"))

        # Invalid paths
        self.assertFalse(is_image_path("document.pdf"))
        self.assertFalse(is_image_path("text.txt"))
        self.assertFalse(is_image_path(""))
        self.assertFalse(is_image_path(None))
        self.assertFalse(is_image_path(123))

    def test_is_photo_question(self):
        """Test is_photo_question function"""
        # Photo-related questions
        self.assertTrue(is_photo_question("Take a photo of the site"))
        self.assertTrue(is_photo_question("Upload image"))
        self.assertTrue(is_photo_question("Picture of equipment"))
        self.assertTrue(is_photo_question("Snapshot of results"))
        self.assertTrue(is_photo_question("Site pic"))

        # Non-photo questions
        self.assertFalse(is_photo_question("What is the temperature?"))
        self.assertFalse(is_photo_question("Enter the pH value"))
        self.assertFalse(is_photo_question("Description of site"))

    @patch('utils.report_generator.STORAGE_PATH', '/test/storage')
    def test_get_full_image_path(self):
        """Test get_full_image_path function"""
        # Test with leading slash
        result = get_full_image_path("/images/test.jpg")
        self.assertEqual(result, "/test/storage/images/test.jpg")

        # Test without leading slash
        result = get_full_image_path("images/test.jpg")
        self.assertEqual(result, "/test/storage/images/test.jpg")

        # Test simple filename
        result = get_full_image_path("test.jpg")
        self.assertEqual(result, "/test/storage/test.jpg")

    @patch('utils.report_generator.os.path.exists')
    @patch('docx.text.run.Run.add_picture')
    def test_add_images_to_cell(self, mock_add_picture, mock_exists):
        """Test add_images_to_cell function"""
        # Mock file existence and path
        mock_exists.return_value = True
        mock_add_picture.return_value = MagicMock()

        # Create mock cell
        mock_cell = MagicMock()
        mock_paragraph = MagicMock()
        mock_run = MagicMock()
        mock_cell.paragraphs = [mock_paragraph]
        mock_paragraph.runs = [mock_run]
        mock_cell.add_paragraph.return_value = mock_paragraph

        # Test adding images
        image_paths = ["/images/test1.jpg", "/images/test2.jpg"]
        add_images_to_cell(mock_cell, image_paths)

        # Verify cell text was cleared
        self.assertEqual(mock_cell.text, "")

        # Verify add_picture was called
        self.assertEqual(mock_run.add_picture.call_count, 2)


class TestBase64ImageSupport(unittest.TestCase):
    """Test suite for base64 image support in report generation"""

    def setUp(self):
        """Set up test fixtures with base64 image data"""
        # Use actual fake signature base64 data (complete and properly padded)
        self.valid_base64_png = (
            "iVBORw0KGgoAAAANSUhEUgAAANIAAAB+CAYAAABRR0/XAAATAUlEQVR4Xu3dC9g"
            "1VVUH8A8LL1RqmtE9pJAMySy7oBEfmhQFQU/1qGQGmqZhdhEppRSyvOYN6KIFfBZd"
            "HgS8IKRm9RJg4S0lSykfeosulhoWqV3V9YPZfPOd773MnDMz75w5az3Pes55z7tn"
            "z8x/9n/2RKAhAohktzcbU5EkUUPgslgiUEcAkc4OPSt0T6g0wimJQCLQEoEyR7Je"
            "lObtluBl8USgIPBp4iQJylyBsyMAAAAASUVORK5CYII="
        )

        # Data URL format
        self.valid_data_url = f"data:image/png;base64,{self.valid_base64_png}"

        # Invalid base64 string
        self.invalid_base64 = "invalid_base64_string"

    def test_is_base64_image_detection(self):
        """Test detection of base64 image data"""
        # Test valid base64 string
        self.assertTrue(is_base64_image(self.valid_base64_png))

        # Test valid data URL
        self.assertTrue(is_base64_image(self.valid_data_url))

        # Test invalid base64
        self.assertFalse(is_base64_image(self.invalid_base64))

        # Test regular file path
        self.assertFalse(is_base64_image("/path/to/image.jpg"))

        # Test non-string input
        self.assertFalse(is_base64_image(123))
        self.assertFalse(is_base64_image(None))

    def test_create_temp_image_from_base64_data_url(self):
        """Test creating temporary image file from data URL"""
        temp_path = create_temp_image_from_base64(self.valid_data_url)

        try:
            # Should create a temporary file
            self.assertIsNotNone(temp_path)
            self.assertTrue(os.path.exists(temp_path))

            # Should have PNG extension
            self.assertTrue(temp_path.endswith('.png'))

            # File should contain image data
            with open(temp_path, 'rb') as f:
                content = f.read()
                # PNG files start with these bytes
                self.assertTrue(content.startswith(b'\x89PNG'))

        finally:
            # Clean up
            if temp_path and os.path.exists(temp_path):
                os.unlink(temp_path)

    def test_create_temp_image_from_base64_plain(self):
        """Test creating temporary image file from plain base64"""
        temp_path = create_temp_image_from_base64(self.valid_base64_png)

        try:
            # Should create a temporary file
            self.assertIsNotNone(temp_path)
            self.assertTrue(os.path.exists(temp_path))

            # Should have PNG extension (default)
            self.assertTrue(temp_path.endswith('.png'))

        finally:
            # Clean up
            if temp_path and os.path.exists(temp_path):
                os.unlink(temp_path)

    def test_create_temp_image_from_invalid_base64(self):
        """Test handling of invalid base64 data"""
        temp_path = create_temp_image_from_base64(self.invalid_base64)

        # Should return None for invalid data
        self.assertIsNone(temp_path)

    def test_get_image_path_or_create_temp_base64(self):
        """Test path resolution for base64 images"""
        path, is_temp = get_image_path_or_create_temp(self.valid_data_url)

        try:
            # Should create temporary file and mark as temp
            self.assertIsNotNone(path)
            self.assertTrue(is_temp)
            self.assertTrue(os.path.exists(path))

        finally:
            # Clean up
            if path and os.path.exists(path):
                os.unlink(path)

    def test_get_image_path_or_create_temp_file_path(self):
        """Test path resolution for regular file paths"""
        file_path = "/path/to/image.jpg"
        path, is_temp = get_image_path_or_create_temp(file_path)

        # Should not create temp file for regular paths
        self.assertFalse(is_temp)
        # Should return full path (with STORAGE_PATH prefix)
        self.assertTrue(path.endswith(file_path))

    @patch('utils.report_generator.os.path.exists')
    @patch('utils.report_generator.os.unlink')
    def test_add_images_to_cell_with_base64(self, mock_unlink, mock_exists):
        """Test adding base64 images to a cell"""
        # Mock document and cell structure
        mock_cell = MagicMock()
        mock_paragraph = MagicMock()
        mock_run = MagicMock()

        mock_cell.paragraphs = [mock_paragraph]
        mock_paragraph.runs = [mock_run]
        mock_cell.add_paragraph.return_value = mock_paragraph
        mock_paragraph.add_run.return_value = mock_run

        # Mock file existence for temp files
        mock_exists.return_value = True

        # Test with base64 image
        image_data = [self.valid_data_url]

        # Call function
        add_images_to_cell(mock_cell, image_data)

        # Verify cell text was cleared
        self.assertEqual(mock_cell.text, "")

        # Verify add_picture was called on the run object
        mock_run.add_picture.assert_called_once()

        # Verify cleanup was called
        mock_unlink.assert_called()

    def test_base64_image_integration_in_report(self):
        """Test integration of base64 images in full report generation"""
        # Create test data with base64 images
        report_data = [
            {
                "name": "Test Group with Images",
                "questions": [
                    {
                        "question": "Photo Documentation",
                        "answers": [self.valid_data_url, self.valid_base64_png]
                    },
                    {
                        "question": "Regular Question",
                        "answers": ["Regular answer"]
                    }
                ]
            }
        ]

        # Create temporary file for report
        with tempfile.NamedTemporaryFile(
            suffix='.docx', delete=False
        ) as temp_file:
            temp_path = temp_file.name

        try:
            # Generate report
            generate_datapoint_report(
                report_data,
                file_path=temp_path,
                form_name="Base64 Image Test Report"
            )

            # Verify report was created
            self.assertTrue(os.path.exists(temp_path))

            # Verify document can be opened
            doc = Document(temp_path)
            self.assertIsNotNone(doc)

            # Should have at least one table
            self.assertGreater(len(doc.tables), 0)

        finally:
            # Clean up
            if os.path.exists(temp_path):
                os.unlink(temp_path)

    def test_mixed_image_types_in_report(self):
        """Test report generation with mixed image types (files and base64)"""
        # Test data with mixed image types
        report_data = [
            {
                "name": "Mixed Image Types",
                "questions": [
                    {
                        "question": "Mixed Images",
                        "answers": [
                            "/path/to/file.jpg",  # File path
                            self.valid_data_url,  # Base64 data URL
                            self.valid_base64_png  # Plain base64
                        ]
                    }
                ]
            }
        ]

        # Create temporary file for report
        with tempfile.NamedTemporaryFile(
            suffix='.docx', delete=False
        ) as temp_file:
            temp_path = temp_file.name

        try:
            # Should not raise any exceptions
            generate_datapoint_report(
                report_data,
                file_path=temp_path,
                form_name="Mixed Image Types Test"
            )

            # Verify report was created
            self.assertTrue(os.path.exists(temp_path))

        finally:
            # Clean up
            if os.path.exists(temp_path):
                os.unlink(temp_path)


if __name__ == '__main__':
    unittest.main()


class TestReportGenerationIntegration(unittest.TestCase):
    """Integration tests that exercise real code paths for better coverage"""

    def setUp(self):
        """Set up test fixtures"""
        self.temp_dir = tempfile.mkdtemp()
        self.test_file_path = os.path.join(self.temp_dir, "test_report.docx")

        # Create a real test image file for integration tests
        self.test_image_path = os.path.join(self.temp_dir, "test_image.png")
        # Create a simple 1x1 pixel PNG
        png_data = base64.b64decode(
            "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAYAAAAfFcSJAAAADUlEQVR42mP"
            "8/5+hHgAHggJ/PchI7wAAAABJRU5ErkJggg=="
        )
        with open(self.test_image_path, 'wb') as f:
            f.write(png_data)

        # Valid base64 data for tests
        self.valid_base64_png = (
            "iVBORw0KGgoAAAANSUhEUgAAANIAAAB+CAYAAABRR0/XAAATAUlEQVR4Xu3dC9g"
            "1VVUH8A8LL1RqmtE9pJAMySy7oBEfmhQFQU/1qGQGmqZhdhEppRSyvOYN6KIFfBZd"
            "HgS8IKRm9RJg4S0lSykfeosulhoWqV3V9YPZfPOd773MnDMz75w5az3Pes55z7tn"
            "z8x/9n/2RKAhAohktzcbU5EkUUPgslgiUEcAkc4OPSt0T6g0wimJQCLQEoEyR7Je"
            "lObtluBl8USgIPBp4iQJylyBsyMAAAAASUVORK5CYII="
        )

    def tearDown(self):
        """Clean up temporary files"""
        import shutil
        shutil.rmtree(self.temp_dir, ignore_errors=True)

    def test_real_image_file_processing(self):
        """Test processing real image files (exercises image code paths)"""
        # Use relative path to test get_full_image_path
        rel_image_path = os.path.relpath(self.test_image_path)

        report_data = [
            {
                "name": "Image Processing Test",
                "questions": [
                    {
                        "question": "Photo of Test Site",
                        "answers": [rel_image_path]
                    }
                ]
            }
        ]

        # This should exercise the actual image processing code
        with patch('utils.report_generator.STORAGE_PATH', self.temp_dir):
            result_path = generate_datapoint_report(
                report_data,
                file_path=self.test_file_path,
                form_name="Real Image Test"
            )

        self.assertTrue(os.path.exists(result_path))
        doc = Document(result_path)
        self.assertGreater(len(doc.tables), 0)

    def test_base64_image_full_integration(self):
        """Test base64 image processing in full report generation"""
        data_url = f"data:image/png;base64,{self.valid_base64_png}"

        report_data = [
            {
                "name": "Base64 Integration Test",
                "questions": [
                    {
                        "question": "Base64 Image Question",
                        "answers": [data_url, self.valid_base64_png]
                    }
                ]
            }
        ]

        # This exercises the base64 image processing code paths
        result_path = generate_datapoint_report(
            report_data,
            file_path=self.test_file_path,
            form_name="Base64 Integration Test"
        )

        self.assertTrue(os.path.exists(result_path))
        doc = Document(result_path)
        self.assertGreater(len(doc.tables), 0)

    def test_mixed_content_integration(self):
        """Test mixed content types in a single report"""
        report_data = [
            {
                "name": "Location Data",
                "questions": [
                    {
                        "question": "Site Name",
                        "answers": ["Site Alpha", "Site Beta", "Site Gamma"]
                    },
                    {
                        "question": "Latitude",
                        "answers": ["12.3456", "12.7890", "12.1234"]
                    },
                    {
                        "question": "Longitude",
                        "answers": ["78.9012", "78.5432", "78.6789"]
                    }
                ]
            },
            {
                "name": "Image Documentation",
                "questions": [
                    {
                        "question": "Site Photo",
                        "answers": [
                            f"data:image/png;base64,{self.valid_base64_png}"
                        ]
                    },
                    {
                        "question": "Equipment Picture",
                        # Test missing file
                        "answers": ["/nonexistent/path.jpg"]
                    }
                ]
            }
        ]

        result_path = generate_datapoint_report(
            report_data,
            file_path=self.test_file_path,
            form_name="Mixed Content Integration Test"
        )

        self.assertTrue(os.path.exists(result_path))
        doc = Document(result_path)
        self.assertGreater(len(doc.tables), 0)

    def test_utility_functions_integration(self):
        """Test utility functions with real data"""
        # Test is_image_path with various inputs
        test_cases = [
            ("/path/image.jpg", True),
            ("photo.PNG", True),
            ("document.pdf", False),
            ("", False),
            (None, False)
        ]

        for input_val, expected in test_cases:
            with self.subTest(input=input_val):
                self.assertEqual(is_image_path(input_val), expected)

        # Test is_photo_question
        photo_questions = [
            "Take a photo of the equipment",
            "Upload image of results",
            "Site picture",
            "Equipment snapshot"
        ]

        non_photo_questions = [
            "Enter the temperature",
            "What is the pH value?",
            "Describe the location"
        ]

        for question in photo_questions:
            with self.subTest(question=question):
                self.assertTrue(is_photo_question(question))

        for question in non_photo_questions:
            with self.subTest(question=question):
                self.assertFalse(is_photo_question(question))

        # Test get_full_image_path
        test_path = "images/test.jpg"
        with patch('utils.report_generator.STORAGE_PATH', '/test/storage'):
            full_path = get_full_image_path(test_path)
            self.assertEqual(full_path, '/test/storage/images/test.jpg')

            # Test with leading slash
            full_path2 = get_full_image_path('/images/test.jpg')
            self.assertEqual(full_path2, '/test/storage/images/test.jpg')

    def test_base64_utility_functions_integration(self):
        """Test base64 utility functions with real data"""
        # Test is_base64_image
        data_url = f"data:image/png;base64,{self.valid_base64_png}"

        self.assertTrue(is_base64_image(data_url))
        self.assertTrue(is_base64_image(self.valid_base64_png))
        self.assertFalse(is_base64_image("not_base64"))
        self.assertFalse(is_base64_image("/path/to/file.jpg"))
        self.assertFalse(is_base64_image(123))

        # Test create_temp_image_from_base64
        temp_path = create_temp_image_from_base64(data_url)
        try:
            self.assertIsNotNone(temp_path)
            self.assertTrue(os.path.exists(temp_path))
            self.assertTrue(temp_path.endswith('.png'))
        finally:
            if temp_path and os.path.exists(temp_path):
                os.unlink(temp_path)

        # Test with invalid data
        invalid_temp = create_temp_image_from_base64("invalid_data")
        self.assertIsNone(invalid_temp)

        # Test get_image_path_or_create_temp
        path, is_temp = get_image_path_or_create_temp(data_url)
        try:
            self.assertTrue(is_temp)
            self.assertIsNotNone(path)
            self.assertTrue(os.path.exists(path))
        finally:
            if path and os.path.exists(path):
                os.unlink(path)

        # Test with regular file path
        file_path = "/path/to/image.jpg"
        path2, is_temp2 = get_image_path_or_create_temp(file_path)
        self.assertFalse(is_temp2)
        self.assertTrue(path2.endswith(file_path))

    def test_error_handling_integration(self):
        """Test error handling scenarios"""
        # Test with invalid file path that will fail during directory creation
        report_data = [
            {
                "name": "Error Test",
                "questions": [
                    {
                        "question": "Test Question",
                        "answers": ["Test Answer"]
                    }
                ]
            }
        ]

        # Mock os.makedirs to force an exception during directory creation
        with patch('utils.report_generator.os.makedirs') as mock_makedirs:
            mock_makedirs.side_effect = PermissionError("Permission denied")
            with self.assertRaises(PermissionError):
                generate_datapoint_report(
                    report_data,
                    file_path="/some/path/test.docx",
                    form_name="Error Test"
                )

    def test_large_dataset_integration(self):
        """Test handling of large datasets that require multiple tables"""
        # Create data with more than 5 answers to trigger table splitting
        large_answers = [f"Answer_{i}" for i in range(8)]

        report_data = [
            {
                "name": "Large Dataset Test",
                "questions": [
                    {
                        "question": "Sample Results",
                        "answers": large_answers
                    },
                    {
                        "question": "Status",
                        "answers": ["Pass"] * 8
                    }
                ]
            }
        ]

        result_path = generate_datapoint_report(
            report_data,
            file_path=self.test_file_path,
            form_name="Large Dataset Test"
        )

        self.assertTrue(os.path.exists(result_path))
        doc = Document(result_path)

        # Should create multiple tables due to >5 answers
        self.assertGreater(len(doc.tables), 1)

    def test_edge_cases_integration(self):
        """Test edge cases and boundary conditions"""
        # Test with None and empty values
        report_data = [
            {
                "name": "Edge Cases",
                "questions": [
                    {
                        "question": "Empty Answers",
                        "answers": []
                    },
                    {
                        "question": "None Values",
                        "answers": [None, None, None]
                    },
                    {
                        "question": "Mixed Empty",
                        "answers": ["Valid", "", None, "Another"]
                    },
                    {
                        "question": "Image with None",
                        "answers": [
                            None,
                            f"data:image/png;base64,{self.valid_base64_png}"
                        ]
                    }
                ]
            }
        ]

        result_path = generate_datapoint_report(
            report_data,
            file_path=self.test_file_path,
            form_name="Edge Cases Test"
        )

        self.assertTrue(os.path.exists(result_path))
        doc = Document(result_path)
        self.assertGreater(len(doc.tables), 0)

    def test_safe_set_cell_text_integration(self):
        """Test safe_set_cell_text with real cell objects"""
        # Create a real document and table for testing
        doc = Document()
        table = doc.add_table(rows=1, cols=3)
        row = table.rows[0]

        # Test valid cell index
        result = safe_set_cell_text(row, 0, "Test Text")
        self.assertTrue(result)
        self.assertEqual(row.cells[0].text, "Test Text")

        # Test None text
        result = safe_set_cell_text(row, 1, None)
        self.assertTrue(result)
        self.assertEqual(row.cells[1].text, "")

        # Test invalid cell index
        result = safe_set_cell_text(row, 10, "Invalid")
        self.assertFalse(result)
